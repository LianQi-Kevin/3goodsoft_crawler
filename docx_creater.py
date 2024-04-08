import csv
import logging
from typing import List, Dict

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Mm, Cm
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from tools.logging_utils import log_set


def load_csv(path: str) -> List[Dict[str, str]]:
    with open(path, 'r', encoding='utf-8') as f:
        return list(csv.DictReader(f))


def init_font_style(paragraph: Paragraph, font_size: float = 4, bold: bool = False,
                    left_indent: float = 0, right_indent: float = 0, space_before: float = 0, space_after: float = 0):
    # 设置字体
    paragraph.style.font.size = Mm(font_size)  # 设置字体大小
    # paragraph.style.font.name = '宋体'  # 设置字体
    paragraph.style.font.bold = bold  # 设置字体加粗
    paragraph.style.font.color.rgb = None  # 设置字体颜色为黑色
    paragraph.style.font.highlight_color = None  # 设置字体背景颜色为无

    # 设置缩进
    paragraph.paragraph_format.left_indent = Mm(left_indent)  # 设置左缩进
    paragraph.paragraph_format.right_indent = Mm(right_indent)  # 设置右缩进

    # 设置段落间距
    paragraph.paragraph_format.space_before = Mm(space_before)  # 设置段前间距
    paragraph.paragraph_format.space_after = Mm(space_after)  # 设置段后间距


def create_docx(path: str, data: List[Dict[str, str]]):
    # docx init
    doc = Document()
    section = doc.sections[0]
    # 设置页面
    section.orientation = WD_ORIENT.PORTRAIT    # 设置纵向
    section.page_width = Mm(210)    # 设置宽度
    section.page_height = Mm(297)   # 设置高度
    # 设置窄边框
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    # 设置中文字体
    doc.styles['Normal'].font.name = u'宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    # add content
    for item in data:
        # set_title
        title_paragraph = doc.add_heading('', level=1)
        init_font_style(title_paragraph, bold=True, font_size=4.5)
        title = title_paragraph.add_run(f"({item['type']}) {item['index']} {item['description']}")
        title.font.name = u'宋体'
        title._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        # set_options
        init_font_style(doc.add_paragraph(f"A. {item['options_A']}"), left_indent=5, space_before=1)
        init_font_style(doc.add_paragraph(f"B. {item['options_B']}"), left_indent=5)
        init_font_style(doc.add_paragraph(f"C. {item['options_C']}"), left_indent=5)
        init_font_style(doc.add_paragraph(f"D. {item['options_D']}"), left_indent=5, space_after=1)
        if item['options_E']:
            init_font_style(doc.add_paragraph(f"E. {item['options_E']}"))

        # set_answer & analysis
        init_font_style(doc.add_paragraph(f"答案：{item['answer']}"), font_size=3)
        init_font_style(doc.add_paragraph(
            f"解析：{item['analysis'] if not item['analysis'].startswith('解析：') else item['analysis'][3:]}"),
                        font_size=3)
        init_font_style(doc.add_paragraph(""))

    # save
    doc.save(path)


if __name__ == '__main__':
    log_set(log_level=logging.DEBUG, log_save=False)
    create_docx("answer_sheet.docx", load_csv("answer_sheet.csv"))
